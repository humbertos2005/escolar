from flask import Blueprint, render_template, request, redirect, url_for, flash, current_app, jsonify, send_file, session
from database import get_db
from .tac_utils import get_next_tac_number
from .utils import login_required, admin_secundario_required
from datetime import datetime
import os
import time

from models_sqlalchemy import (
    TAC, TACObrigacao, TACParticipante, Aluno, Cabecalho, DadosEscola
)

# Optional imports for DOCX export
from docx import Document
import io

formularios_tac_bp = Blueprint('formularios_tac_bp', __name__)

@formularios_tac_bp.route('/tacs')
@admin_secundario_required
def listar_tacs():
    db = get_db()
    show_deleted = request.args.get('show_deleted') == '1'
    try:
        query = db.query(TAC)
        if not show_deleted:
            query = query.filter_by(deleted='0')
        tacs = []
        for t in query.order_by(TAC.created_at.desc()).all():
            tac = t.__dict__.copy()
            tac['numero_display'] = tac.get('numero') or ''
            tac['aluno_nome'] = None
            tac['serie_turma'] = None
            if t.aluno_id:
                aluno = db.query(Aluno).filter_by(id=t.aluno_id).first()
                if aluno:
                    tac['aluno_nome'] = aluno.nome
                    tac['serie_turma'] = f"{aluno.serie or ''} {aluno.turma or ''}".strip()
            if not tac.get('aluno_nome'):
                tac['aluno_nome'] = tac.get('escola_text') or '-'
            tacs.append(tac)
        return render_template('formularios/listar_tacs.html', tacs=tacs, show_deleted=show_deleted)
    except Exception:
        current_app.logger.exception("Erro ao listar TACS")
        flash('Erro ao listar TACS', 'danger')
        return render_template('formularios/listar_tacs.html', tacs=[], show_deleted=show_deleted)

@formularios_tac_bp.route('/tac/novo', methods=['GET', 'POST'])
@admin_secundario_required
def tac_novo():
    db = get_db()
    if request.method == 'POST':
        aluno_id = request.form.get('aluno_id') or None
        try:
            aluno_id = int(aluno_id) if aluno_id else None
        except Exception:
            aluno_id = None
        cabecalho_id = request.form.get('cabecalho_id') or None
        escola_text = request.form.get('escola','').strip()
        serie = request.form.get('serie','').strip()
        turma = request.form.get('turma','').strip()
        responsavel = request.form.get('responsavel','').strip()
        diretor = request.form.get('diretor','').strip()
        fato = request.form.get('fato','').strip()
        prazo = request.form.get('prazo','').strip()
        obrigacoes = request.form.getlist('obrigacao[]')
        participantes_nomes = request.form.getlist('participante_nome[]')
        participantes_cargos = request.form.getlist('participante_cargo[]')

        max_attempts = 8
        attempt = 0
        tac_obj = None
        while attempt < max_attempts:
            attempt += 1
            try:
                year = datetime.utcnow().strftime('%Y')
                base_count = db.query(TAC).filter(TAC.created_at.like(f"{year}%")).count()
                seq = base_count + attempt
                numero = f"TAC-{seq:04d}/{year}"
                tac_obj = TAC(
                    numero=numero,
                    aluno_id=aluno_id,
                    cabecalho_id=cabecalho_id,
                    escola_text=escola_text,
                    serie=serie,
                    turma=turma,
                    responsavel=responsavel,
                    diretor_nome=diretor,
                    fato=fato,
                    prazo=prazo,
                    created_at=datetime.utcnow().isoformat(),
                    updated_at=datetime.utcnow().isoformat(),
                    deleted='0'
                )
                db.add(tac_obj)
                db.commit()
                # inserir obrigações
                for idx, ob in enumerate([o for o in obrigacoes if o and o.strip()]):
                    db.add(TACObrigacao(tac_id=tac_obj.id, descricao=ob.strip(), ordem=idx))
                # inserir participantes
                for idx, (pn, pc) in enumerate(zip(participantes_nomes, participantes_cargos)):
                    if pn and pn.strip():
                        db.add(TACParticipante(tac_id=tac_obj.id, nome=pn.strip(), cargo=(pc or '').strip(), ordem=idx))
                db.commit()
                flash('TAC criado com sucesso.', 'success')
                return redirect(url_for('formularios_tac_bp.listar_tacs'))
            except Exception as e:
                db.rollback()
                if 'UNIQUE constraint failed: tacs.numero' in str(e):
                    time.sleep(0.05)
                    continue
                current_app.logger.exception("Erro ao salvar TAC.")
                flash('Erro ao salvar TAC.', 'danger')
                return redirect(url_for('formularios_tac_bp.tac_novo'))

        # Fallback: após várias tentativas, gerar um número baseado em timestamp
        try:
            fallback_seq = datetime.utcnow().strftime('%Y%m%d%H%M%S%f')[-8:]
            fallback_num = f"TAC-{fallback_seq}/{datetime.utcnow().year}"
            current_app.logger.info(f"Usando número-fallback para TAC: {fallback_num}")
            tac_obj = TAC(
                numero=fallback_num,
                aluno_id=aluno_id,
                cabecalho_id=cabecalho_id,
                escola_text=escola_text,
                serie=serie,
                turma=turma,
                responsavel=responsavel,
                diretor_nome=diretor,
                fato=fato,
                prazo=prazo,
                created_at=datetime.utcnow().isoformat(),
                updated_at=datetime.utcnow().isoformat(),
                deleted='0'
            )
            db.add(tac_obj)
            db.commit()
            for idx, ob in enumerate([o for o in obrigacoes if o and o.strip()]):
                db.add(TACObrigacao(tac_id=tac_obj.id, descricao=ob.strip(), ordem=idx))
            for idx, (pn, pc) in enumerate(zip(participantes_nomes, participantes_cargos)):
                if pn and pn.strip():
                    db.add(TACParticipante(tac_id=tac_obj.id, nome=pn.strip(), cargo=(pc or '').strip(), ordem=idx))
            db.commit()
            flash('TAC criado com sucesso.', 'success')
            return redirect(url_for('formularios_tac_bp.listar_tacs'))
        except Exception:
            db.rollback()
            current_app.logger.exception("Falha ao inserir TAC mesmo com numero-fallback.")
            flash('Erro ao salvar TAC. Falha ao gerar número único.', 'danger')
            return redirect(url_for('formularios_tac_bp.tac_novo'))

    return render_template('formularios/tac_form.html', tac=None)

@formularios_tac_bp.route('/tac/editar/<int:id>', methods=['GET', 'POST'])
@admin_secundario_required
def tac_editar(id):
    db = get_db()
    t = db.query(TAC).filter_by(id=id).first()
    if not t:
        flash('TAC não encontrado.', 'warning')
        return redirect(url_for('formularios_tac_bp.listar_tacs'))
    tac = t.__dict__.copy()
    if request.method == 'POST':
        aluno_id = request.form.get('aluno_id') or None
        try:
            aluno_id = int(aluno_id) if aluno_id else None
        except Exception:
            aluno_id = None
        cabecalho_id = request.form.get('cabecalho_id') or None
        escola_text = request.form.get('escola','').strip()
        serie = request.form.get('serie','').strip()
        turma = request.form.get('turma','').strip()
        responsavel = request.form.get('responsavel','').strip()
        diretor = request.form.get('diretor','').strip()
        fato = request.form.get('fato','').strip()
        prazo = request.form.get('prazo','').strip()

        obrigacoes = request.form.getlist('obrigacao[]')
        participantes_nomes = request.form.getlist('participante_nome[]')
        participantes_cargos = request.form.getlist('participante_cargo[]')

        try:
            t.aluno_id = aluno_id
            t.cabecalho_id = cabecalho_id
            t.escola_text = escola_text
            t.serie = serie
            t.turma = turma
            t.responsavel = responsavel
            t.diretor_nome = diretor
            t.fato = fato
            t.prazo = prazo
            t.updated_at = datetime.utcnow().isoformat()
            db.commit()
            db.query(TACObrigacao).filter_by(tac_id=id).delete()
            for idx, ob in enumerate([o for o in obrigacoes if o and o.strip()]):
                db.add(TACObrigacao(tac_id=id, descricao=ob.strip(), ordem=idx))
            db.query(TACParticipante).filter_by(tac_id=id).delete()
            for idx, (pn, pc) in enumerate(zip(participantes_nomes, participantes_cargos)):
                if pn and pn.strip():
                    db.add(TACParticipante(tac_id=id, nome=pn.strip(), cargo=(pc or '').strip(), ordem=idx))
            db.commit()
            flash('TAC atualizado com sucesso.', 'success')
            return redirect(url_for('formularios_tac_bp.listar_tacs'))
        except Exception:
            current_app.logger.exception("Erro ao atualizar TAC")
            db.rollback()
            flash('Erro ao atualizar TAC.', 'danger')
            return redirect(url_for('formularios_tac_bp.tac_editar', id=id))

    obrig = db.query(TACObrigacao).filter_by(tac_id=id).order_by(TACObrigacao.ordem).all()
    part = db.query(TACParticipante).filter_by(tac_id=id).order_by(TACParticipante.ordem).all()
    tac['obrigacoes'] = [o.descricao for o in obrig]
    tac['participantes'] = [{'nome': p.nome, 'cargo': p.cargo} for p in part]
    return render_template('formularios/tac_form.html', tac=tac)

@formularios_tac_bp.route('/tac/visualizar/<int:id>')
@admin_secundario_required
def tac_visualizar(id):
    db = get_db()
    t = db.query(TAC).filter_by(id=id).first()
    if not t:
        flash('TAC não encontrado.', 'warning')
        return redirect(url_for('formularios_tac_bp.listar_tacs'))
    tac = t.__dict__.copy()
    tac['aluno'] = None
    if t.aluno_id:
        a = db.query(Aluno).filter_by(id=t.aluno_id).first()
        tac['aluno'] = a.__dict__ if a else None
    tac['obrigacoes'] = [o.descricao for o in db.query(TACObrigacao).filter_by(tac_id=id).order_by(TACObrigacao.ordem).all()]
    tac['participantes'] = [ {'nome': p.nome, 'cargo': p.cargo} for p in db.query(TACParticipante).filter_by(tac_id=id).order_by(TACParticipante.ordem).all()]
    # Derivados p/ template
    try:
        serie = tac.get('serie') or ''
        turma = tac.get('turma') or ''
        if serie and turma:
            tac['serie_turma_full'] = f"{serie} / {turma}"
        elif serie:
            tac['serie_turma_full'] = serie
        else:
            tac['serie_turma_full'] = tac.get('serie_turma') or ''
    except Exception:
        tac['serie_turma_full'] = tac.get('serie_turma') or ''
    try:
        aluno_nome = ''
        if tac.get('aluno'):
            aluno_nome = tac['aluno'].get('nome') or ''
        aluno_nome = aluno_nome or tac.get('aluno_nome') or ''
        tac['aluno_nome_upper'] = aluno_nome.upper() if aluno_nome else ''
    except Exception:
        tac['aluno_nome_upper'] = (tac.get('aluno_nome') or '').upper()
    # Cabeçalho
    cabecalho = None
    try:
        cab_id = tac.get('cabecalho_id') or None
        if cab_id:
            ch = db.query(Cabecalho).filter_by(id=cab_id).first()
            if ch:
                cabecalho = ch.__dict__.copy()
                for field in ['logo_escola', 'logo_secretaria', 'logo_estado']:
                    fn = getattr(ch, field, None)
                    if fn:
                        cabecalho[f"{field}_url"] = url_for('static', filename=f'uploads/cabecalhos/{os.path.basename(fn)}') if os.path.exists(os.path.join(current_app.root_path, 'static', 'uploads', 'cabecalhos', os.path.basename(fn))) else None
                de = db.query(DadosEscola).filter_by(cabecalho_id=cab_id).first()
                cabecalho['dados_escola'] = de.__dict__ if de else None
    except Exception:
        current_app.logger.exception("Erro ao buscar cabecalho para TAC visualizar")
        cabecalho = None

    return render_template('formularios/tac_view.html', tac=tac, cabecalho=cabecalho)

@formularios_tac_bp.route('/tac/excluir/<int:id>', methods=['POST'])
@admin_secundario_required
def tac_excluir(id):
    db = get_db()
    try:
        t = db.query(TAC).filter_by(id=id).first()
        if t:
            t.deleted = '1'
            t.updated_at = datetime.utcnow().isoformat()
            db.commit()
            flash('TAC excluído (soft-delete).', 'success')
    except Exception:
        current_app.logger.exception("Erro ao excluir TAC")
        db.rollback()
        flash('Erro ao excluir TAC.', 'danger')
    return redirect(url_for('formularios_tac_bp.listar_tacs'))

@formularios_tac_bp.route('/api/alunos_autocomplete')
def alunos_autocomplete():
    q = request.args.get('q','').strip()
    db = get_db()
    query = db.query(Aluno)
    if q and len(q) > 0:
        ilike = f"%{q.upper()}%"
        alunos = query.filter(Aluno.nome.ilike(ilike)).order_by(Aluno.nome).limit(30).all()
    else:
        alunos = query.order_by(Aluno.id.desc()).limit(30).all()
    return jsonify([{'id': a.id, 'nome': a.nome} for a in alunos])

@formularios_tac_bp.route('/api/aluno')
def api_aluno():
    aid = request.args.get('id')
    db = get_db()
    if not aid:
        return jsonify({'error':'missing id'}), 400
    try:
        a = db.query(Aluno).filter_by(id=aid).first()
        if not a:
            return jsonify({'error':'not found'}), 404
        # Retorna apenas os campos básicos do aluno
        return jsonify({
            'id': a.id,
            'nome': a.nome,
            'serie': a.serie,
            'turma': a.turma,
            'responsavel': a.responsavel
        })
    except Exception as e:
        current_app.logger.exception('Erro ao buscar aluno por id')
        return jsonify({'error': 'internal', 'detail': str(e)}), 500

# EXPORT DOCX (OPCIONAL)
@formularios_tac_bp.route('/tac/<int:id>/export_docx')
@login_required
def export_docx(id):
    db = get_db()
    t = db.query(TAC).filter_by(id=id).first()
    if not t:
        flash('TAC não encontrado.', 'warning')
        return redirect(url_for('formularios_tac_bp.listar_tacs'))
    tac = t.__dict__.copy()
    tac['aluno'] = None
    if t.aluno_id:
        a = db.query(Aluno).filter_by(id=t.aluno_id).first()
        tac['aluno'] = a.__dict__ if a else None
    tac['obrigacoes'] = [o.descricao for o in db.query(TACObrigacao).filter_by(tac_id=id).order_by(TACObrigacao.ordem).all()]
    tac['participantes'] = [ {'nome': p.nome, 'cargo': p.cargo} for p in db.query(TACParticipante).filter_by(tac_id=id).order_by(TACParticipante.ordem).all() ]

    doc = Document()
    doc.add_heading('TERMO DE ADEQUAÇÃO DE CONDUTA (TAC)', level=1)
    doc.add_paragraph(f'Número: {tac.get("numero","-")}')
    doc.add_paragraph(f'Escola: {tac.get("escola_text","-")}')
    aluno_nome = tac.get('aluno',{}).get('nome') if tac.get('aluno') else tac.get('aluno_nome')
    doc.add_paragraph(f'Aluno: {aluno_nome or "-"}')
    doc.add_paragraph(f'Série / Turma: {tac.get("serie","-")} / {tac.get("turma","-")}')
    doc.add_paragraph(f'Responsável: {tac.get("responsavel","-")}')
    doc.add_paragraph(f'Diretor: {tac.get("diretor_nome","-")}')
    doc.add_paragraph('---')
    doc.add_heading('Do Fato', level=2)
    doc.add_paragraph(tac.get('fato',''))
    doc.add_heading('Das Obrigações e Prazo', level=2)
    if tac.get('obrigacoes'):
        for ob in tac['obrigacoes']:
            doc.add_paragraph(ob, style='List Number')
    else:
        doc.add_paragraph('Sem obrigações registradas.')
    doc.add_heading('Participantes', level=2)
    if tac.get('participantes'):
        for p in tac['participantes']:
            doc.add_paragraph(f"{p.get('nome','-')} — {p.get('cargo','')}")
    else:
        doc.add_paragraph('Sem participantes registrados.')
    diretor_text = tac.get('diretor_nome') or 'Diretor(a) da Escola'
    doc.add_paragraph('\n\n\n__________________________\n' + f"Diretor(a): {diretor_text}")
    if tac.get('participantes'):
        for p in tac['participantes']:
            nome = p.get('nome') or '-'
            cargo = p.get('cargo') or ''
            doc.add_paragraph('\n\n\n__________________________\n' + f"{nome} — {cargo}")
    else:
        resp = tac.get('responsavel') or '-'
        doc.add_paragraph('\n\n\n__________________________\n' + f"Responsável: {resp}")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    filename = f"TAC-{tac.get('numero','id'+str(id))}.docx"
    return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# EXPORT DOCX by docxtpl
from docxtpl import DocxTemplate

@formularios_tac_bp.route('/tac/<int:id>/export_docx_template')
@login_required
def export_docx_template_docxtpl(id):
    db = get_db()
    t = db.query(TAC).filter_by(id=id).first()
    if not t:
        flash('TAC não encontrado.', 'warning')
        return redirect(url_for('formularios_tac_bp.listar_tacs'))
    tac = t.__dict__.copy()
    tac['aluno'] = None
    if t.aluno_id:
        a = db.query(Aluno).filter_by(id=t.aluno_id).first()
        tac['aluno'] = a.__dict__ if a else None
    obrigacoes = [o.descricao for o in db.query(TACObrigacao).filter_by(tac_id=id).order_by(TACObrigacao.ordem).all()]
    participantes = [ {'nome': p.nome, 'cargo': p.cargo} for p in db.query(TACParticipante).filter_by(tac_id=id).order_by(TACParticipante.ordem).all() ]

    template_path = current_app.config.get('TAC_DOCX_TEMPLATE_PATH', 'templates/docx/tac_template.docx')
    try:
        doc = DocxTemplate(template_path)
        context = {
            'numero': tac.get('numero', ''),
            'escola': tac.get('escola_text') or tac.get('cabecalho_nome', ''),
            'aluno_nome': (tac.get('aluno') and tac['aluno'].get('nome')) or tac.get('aluno_nome',''),
            'serie': tac.get('serie',''),
            'turma': tac.get('turma',''),
            'responsavel': tac.get('responsavel',''),
            'diretor': tac.get('diretor_nome',''),
            'fato': tac.get('fato',''),
            'prazo': tac.get('prazo',''),
            'obrigacoes': obrigacoes,
            'participantes': participantes,
            'data': (tac.get('updated_at') or tac.get('created_at') or '')[:10],
            'emitido_por': (session.get('username') or '')
        }
        bio = io.BytesIO()
        doc.render(context)
        doc.save(bio)
        bio.seek(0)
        filename = f"TAC-{tac.get('numero', 'id'+str(id))}.docx"
        return send_file(
            bio,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception:
        current_app.logger.exception("Erro ao gerar DOCX do TAC")
        flash('Erro ao gerar DOCX do TAC. Verifique o template e os placeholders.', 'danger')
        return redirect(url_for('formularios_tac_bp.tac_visualizar', id=id))
